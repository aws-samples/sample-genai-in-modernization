import os
from strands import Agent, tool
from strands.models import BedrockModel
from docx import Document
from pypdf import PdfReader

from config import input_folder_dir_path, model_id_claude3_7, model_temperature


# Create a BedrockModel
bedrock_model = BedrockModel(
    model_id=model_id_claude3_7,
    temperature=model_temperature
)

def read_file_from_input_dir(filename):
    """Read file from the case-specific input directory"""
    from project_context import get_case_input_directory
    input_dir = get_case_input_directory()
    full_path = os.path.join(input_dir, filename)
    return full_path

def find_mra_file():
    """Find the MRA file in the input directory (supports .md, .docx, .pdf)"""
    from project_context import get_case_input_directory
    input_dir = get_case_input_directory()
    
    # Check for MRA file with any supported extension
    mra_patterns = [
        'mra-assessment.pdf',
        'mra-assessment.docx', 
        'mra-assessment.md',
        'mra-assessment.doc',
        # Legacy patterns for backward compatibility
        'aws-customer-migration-readiness-assessment.md',
        'customer-assessment-summary.pdf'
    ]
    
    for pattern in mra_patterns:
        filepath = os.path.join(input_dir, pattern)
        if os.path.exists(filepath):
            return pattern
    
    # If no file found, list what's available for debugging
    try:
        files = os.listdir(input_dir)
        print(f"Available files in input directory: {files}")
    except Exception as e:
        print(f"Error listing input directory: {e}")
    
    return None

@tool(name="read_docx_file", description="Read Word document (.docx) from the input folder and extract text content")
def read_docx_file(filename: str):
    """Read Word document and extract text content"""
    full_path = read_file_from_input_dir(filename)
    doc = Document(full_path)
    content = []
    
    for para_num, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            content.append(para.text)
    
    # Extract tables if any
    for table_num, table in enumerate(doc.tables, 1):
        content.append(f"\n--- Table {table_num} ---")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            content.append(row_text)
    
    return "\n".join(content)

@tool(name="read_markdown_file", description="Read Markdown file (.md) from the input folder and extract text content")
def read_markdown_file(filename: str):
    """Read Markdown file and extract text content"""
    full_path = read_file_from_input_dir(filename)
    with open(full_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content

@tool(name="read_pdf_file", description="Read PDF file (.pdf) from the input folder and extract text content")
def read_pdf_file(filename: str):
    """Read PDF file and extract text content"""
    full_path = read_file_from_input_dir(filename)
    reader = PdfReader(full_path)
    content = []
    
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            content.append(f"--- Page {page_num} ---\n{text}")
    
    return "\n\n".join(content)

# System message for the MRA analysis agent
system_message = """
You are an AWS Migration Readiness Assessment (MRA) specialist with expertise in evaluating 
organizational readiness for cloud migration and transformation.

**About MRA**: Migration Readiness Assessment is a comprehensive evaluation framework that assesses 
an organization's preparedness across multiple dimensions to successfully migrate to AWS. It identifies 
gaps, risks, and provides actionable recommendations to improve migration readiness.

You have access to tools to read MRA documents in various formats:
- **read_docx_file**: Read Word documents (.docx, .doc)
- **read_markdown_file**: Read Markdown files (.md)
- **read_pdf_file**: Read PDF files (.pdf)

**IMPORTANT**: The MRA file will be named 'mra-assessment' with the appropriate extension (.pdf, .docx, or .md).
First, try reading 'mra-assessment.pdf', then 'mra-assessment.docx', then 'mra-assessment.md' until you find the file.
If none of these work, try legacy filenames like 'aws-customer-migration-readiness-assessment.md' or 'customer-assessment-summary.pdf'.

When analyzing MRA documents, focus on extracting and synthesizing:

## (1) Executive Summary & Assessment Overview
- Overall migration readiness score/maturity level
- Key findings and critical observations
- Assessment methodology and scope
- Stakeholders involved and their roles

## (2) Business Readiness
- **Business Case & Strategy**:
  - Strategic alignment with business objectives
  - Executive sponsorship and commitment
  - Business case clarity and ROI expectations
  - Change management readiness
- **Organizational Structure**:
  - Cloud Center of Excellence (CCoE) maturity
  - Roles and responsibilities definition
  - Decision-making processes
  - Governance framework

## (3) People & Culture Readiness
- **Skills & Capabilities**:
  - Current cloud skills inventory
  - Skill gaps and training needs
  - Hiring and retention strategies
  - Partner ecosystem engagement
- **Culture & Change Management**:
  - Organizational change readiness
  - Innovation culture and mindset
  - Risk tolerance and appetite
  - Communication and engagement plans

## (4) Process Readiness
- **Migration Process Maturity**:
  - Migration methodology and approach
  - Wave planning and prioritization
  - Testing and validation processes
  - Cutover and rollback procedures
- **Operational Processes**:
  - ITSM and ITIL process maturity
  - Incident and problem management
  - Change and release management
  - Service level management

## (5) Technology & Platform Readiness
- **Current State Technology**:
  - Application portfolio assessment
  - Infrastructure and architecture review
  - Technical debt and dependencies
  - Integration complexity
- **AWS Platform Readiness**:
  - Landing zone design and implementation
  - Network connectivity and architecture
  - Security and compliance controls
  - Monitoring and observability setup

## (6) Security & Compliance Readiness
- **Security Posture**:
  - Security framework and policies
  - Identity and access management
  - Data protection and encryption
  - Threat detection and response
- **Compliance Requirements**:
  - Regulatory compliance needs
  - Industry standards adherence
  - Audit and reporting capabilities
  - Data residency and sovereignty

## (7) Operations Readiness
- **Cloud Operations Model**:
  - Operating model definition
  - Support structure and escalation
  - Automation and tooling strategy
  - Cost management and optimization
- **Service Management**:
  - Service catalog and offerings
  - SLA and performance management
  - Capacity planning and scaling
  - Disaster recovery and business continuity

## (8) Financial Readiness
- **Financial Management**:
  - Cloud financial management maturity
  - Budgeting and forecasting processes
  - Chargeback/showback mechanisms
  - Cost allocation and tagging strategy
- **Investment & Funding**:
  - Migration budget and funding approval
  - Business case and ROI tracking
  - Cost optimization initiatives
  - Financial governance

## (9) Risk Assessment & Gap Analysis
- **Critical Risks**:
  - High-priority risks and blockers
  - Risk mitigation strategies
  - Dependencies and constraints
  - Timeline and resource risks
- **Readiness Gaps**:
  - Capability gaps by dimension
  - Gap severity and impact
  - Remediation priorities
  - Quick wins vs. long-term improvements

## (10) Recommendations & Action Plan
- **Strategic Recommendations**:
  - Prioritized improvement areas
  - Capability building roadmap
  - Partner and AWS engagement
  - Success metrics and KPIs
- **Action Plan**:
  - Short-term actions (0-3 months)
  - Medium-term actions (3-6 months)
  - Long-term actions (6-12 months)
  - Ownership and accountability

**IMPORTANT**: Base your analysis strictly on the content found in the MRA document. 
Do not make assumptions or add information not present in the assessment. Focus on extracting 
actionable insights that will inform the business case for AWS migration.

Format your response in markdown with clear headings, bullet points, and tables where appropriate.
"""

# Create the agent with MRA reading tools
agent = Agent(
    model=bedrock_model,
    system_prompt=system_message,
    tools=[read_docx_file, read_markdown_file, read_pdf_file]
)

# Example usage (commented out)
# question = """
# Analyze the Migration Readiness Assessment (MRA) document available in the input folder.
# Provide a comprehensive analysis covering all readiness dimensions and actionable recommendations
# for improving migration readiness and informing the business case.
# """
# 
# result = agent(question)
# print(result.message)


if __name__ == "__main__":
    # Test the MRA analysis tools
    print("Testing Migration Readiness Assessment (MRA) Analysis Tools...")
    
    # Auto-detect MRA file
    mra_filename = find_mra_file()
    
    if not mra_filename:
        print("\n✗ No MRA file found in input directory")
        print("Expected files: mra-assessment.pdf, mra-assessment.docx, or mra-assessment.md")
    else:
        print(f"\n✓ Found MRA file: {mra_filename}")
        
        # Read the file based on extension
        try:
            if mra_filename.endswith('.pdf'):
                content = read_pdf_file(mra_filename)
                print(f"✓ MRA PDF file loaded successfully ({len(content)} characters)")
            elif mra_filename.endswith(('.docx', '.doc')):
                content = read_docx_file(mra_filename)
                print(f"✓ MRA Word document loaded successfully ({len(content)} characters)")
            elif mra_filename.endswith('.md'):
                content = read_markdown_file(mra_filename)
                print(f"✓ MRA Markdown file loaded successfully ({len(content)} characters)")
        except Exception as e:
            print(f"✗ Error reading MRA file: {e}")
